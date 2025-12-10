import json
import os
import re
import shutil
from io import BytesIO
from datetime import datetime
from docx import Document
from config import PASTA_MEMORIA, ARQUIVO_CONFIG

def carregar_config_json():
    if not os.path.exists(PASTA_MEMORIA): os.makedirs(PASTA_MEMORIA)
    if not os.path.exists(ARQUIVO_CONFIG):
        with open(ARQUIVO_CONFIG, 'w') as f: json.dump({"pinned": [], "trash": []}, f)
    try:
        with open(ARQUIVO_CONFIG, 'r') as f:
            cfg = json.load(f)
            if "trash" not in cfg: cfg["trash"] = []
            return cfg
    except: return {"pinned": [], "trash": []}

def salvar_config_json(config):
    with open(ARQUIVO_CONFIG, 'w') as f: json.dump(config, f)

def faxina_inicial():
    """Tenta apagar pastas marcadas como lixo"""
    cfg = carregar_config_json()
    lixo = cfg.get("trash", [])
    if not lixo: return
    
    restante = []
    for caso in lixo:
        caminho = os.path.join(PASTA_MEMORIA, caso)
        if os.path.exists(caminho):
            try: shutil.rmtree(caminho) 
            except: restante.append(caso) 
            
    if len(restante) != len(lixo):
        cfg["trash"] = restante
        salvar_config_json(cfg)

def gerar_word(texto):
    doc = Document()
    doc.add_heading('Relatório Nemesis AI', 0)
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("-" * 50)
    doc.add_paragraph(texto)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def listar_casos_visiveis():
    if not os.path.exists(PASTA_MEMORIA): os.makedirs(PASTA_MEMORIA)
    cfg = carregar_config_json()
    todos = [f for f in os.listdir(PASTA_MEMORIA) 
             if os.path.isdir(os.path.join(PASTA_MEMORIA, f)) 
             and not f.startswith("_") 
             and f not in cfg.get("trash", [])]
    pinned = [c for c in todos if c in cfg.get("pinned", [])]
    others = [c for c in todos if c not in cfg.get("pinned", [])]
    return sorted(pinned) + sorted(others)